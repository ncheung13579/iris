"""Build IRIS_Final_Report_Nathan_Cheung.{docx,pdf} from docs/Project_Report.md.

Pipeline:
  1. pandoc   docs/Project_Report.md  ->  .docx   (with TOC, figures, YAML title page)
  2. python-docx post-process          ->  justify body paragraphs
                                           (code, tables, headings, captions untouched)
  3. docx2pdf (via Word)               ->  .pdf

Run from repo root:
    python scripts/build_report.py
"""
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from pypdf import PdfReader

ROOT = Path(__file__).parent.parent
MD   = ROOT / "docs" / "Project_Report.md"
DOCX = ROOT / "IRIS_Final_Report_Nathan_Cheung.docx"
PDF  = ROOT / "IRIS_Final_Report_Nathan_Cheung.pdf"

# Paragraphs that should be justified. Everything else (headings, code, title-
# page elements, figure captions, table cells) stays at its pandoc default.
BODY_STYLES = {"Body Text", "First Paragraph", "Compact", "Normal"}


def step_1_pandoc() -> None:
    pypandoc.convert_file(
        str(MD), "docx", outputfile=str(DOCX),
        extra_args=[
            "--toc", "--toc-depth=2",
            "--standalone",
            f"--resource-path={ROOT / 'docs'};{ROOT}",
        ],
    )


def step_2_justify() -> int:
    """Justify body paragraphs in-place. Returns count justified.

    python-docx's `doc.paragraphs` only enumerates top-level paragraphs — not
    those inside table cells — so tables are automatically excluded. Code
    blocks carry the 'Source Code' style which is not in BODY_STYLES, so they
    are also left alone.
    """
    doc = Document(str(DOCX))
    n = 0
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        if style in BODY_STYLES:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            n += 1
    doc.save(str(DOCX))
    return n


def step_3_pdf() -> None:
    convert(str(DOCX), str(PDF))


def main() -> None:
    print("[1/3] pandoc markdown -> docx ...")
    step_1_pandoc()
    print(f"      wrote {DOCX.name}  ({DOCX.stat().st_size:,} bytes)")

    print("[2/3] justifying body paragraphs ...")
    n = step_2_justify()
    print(f"      justified {n} paragraphs (code/tables/headings/captions untouched)")

    print("[3/3] docx -> pdf (via Word) ...")
    step_3_pdf()
    pages = len(PdfReader(str(PDF)).pages)
    print(f"      wrote {PDF.name}  ({PDF.stat().st_size:,} bytes, {pages} pages)")


if __name__ == "__main__":
    main()
